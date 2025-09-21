from flask import Blueprint, render_template, session, redirect, url_for, flash, request, jsonify,send_file
from psycopg2 import extras
from datetime import date
from conexion import get_connection
import base64
from io import BytesIO
from reportlab.lib.pagesizes import  A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import cm
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openpyxl import Workbook
from tempfile import NamedTemporaryFile
import tempfile
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
import os
from openpyxl.drawing.image import Image as XLImage
from openpyxl.styles import Font, Alignment
from io import BytesIO

caracequipos_bp = Blueprint("caracequipos", __name__, template_folder="templates")

# 🔹 Vista principal
@caracequipos_bp.route("/equiposcaracteristicas/<int:id_equipo>")
def equiposcaracteristicas(id_equipo):
    if "usuario" not in session:
        flash("Debes iniciar sesión primero", "warning")
        return redirect(url_for("inventarioequipos"))

    conn = get_connection()
    cur = conn.cursor(cursor_factory=extras.RealDictCursor)

    # Equipo
    cur.execute("""
        SELECT e.*, s.sede, o.oficina_salon, p.piso, u.nombre as usuario_nombre, u.cc as usuario_cc, u.email as usuario_email,ar.area as area,ca.descripcion as cargo
        FROM equipo e
        LEFT JOIN sede s ON e.idsede = s.id
        LEFT JOIN oficina o ON e.idoficina = o.id
        LEFT JOIN piso p ON e.idpiso = p.id
        LEFT JOIN usuario u ON e.idusuario = u.id
        LEFT JOIN cargo ca ON ca.id = u.idcargo 
        LEFT JOIN area ar ON ar.id = u.idarea             
        WHERE e.id = %s
    """, (id_equipo,))
    equipo = cur.fetchone()

    # Características
    cur.execute("""
        SELECT c.*, b.tarjetamadre,b.chipset, so.descripcion as sistema,so.arquitectura, pr.modelo as procesador,pr.marca,pr.frecuencia,pr.nucleos,pr.plataforma,
               ud.nombre as usuario_dominio,idsistemaoperativo,idboard,idprocesador,idusuariodominio
        FROM caracteristicasequipo c
        LEFT JOIN board b ON c.idboard = b.id
        LEFT JOIN sistemaopera so ON c.idsistemaoperativo = so.id
        LEFT JOIN procesador pr ON c.idprocesador = pr.id
        LEFT JOIN usuariodominio ud ON c.idusuariodominio = ud.id
        WHERE c.idequipo = %s
    """, (id_equipo,))
    caracteristicas = cur.fetchone()

    # Datos para los selects
    cur.execute("SELECT id, tarjetamadre FROM board WHERE baja = FALSE ORDER BY tarjetamadre")
    boards = cur.fetchall()

    cur.execute("SELECT id, descripcion FROM sistemaopera ORDER BY descripcion")
    sistemas = cur.fetchall()

    cur.execute("SELECT id, modelo FROM procesador WHERE baja = FALSE ORDER BY modelo")
    procesadores = cur.fetchall()

    cur.execute("SELECT id, nombre FROM usuariodominio ORDER BY nombre")
    usuarios_dominio = cur.fetchall()
        # Datos para los selects extra (equipo)
    cur.execute("SELECT id, sede FROM sede ORDER BY sede")
    sedes = cur.fetchall()

    cur.execute("SELECT id, oficina_salon FROM oficina ORDER BY oficina_salon")
    oficinas = cur.fetchall()

    cur.execute("SELECT id, piso FROM piso ORDER BY piso")
    pisos = cur.fetchall()

    cur.execute("SELECT id, nombre FROM usuario WHERE baja = FALSE ORDER BY nombre")
    usuarios = cur.fetchall()


   # 🔹 Almacenamiento asignado
    cur.execute("""
        SELECT a.id, a.marca, a.horas_uso, a.estado, c.capacidad, c.tipo
        FROM almacenamiento a
        JOIN almacenamientocarac c ON a.idalmacenamientocarac = c.id
        WHERE a.idequipo = %s
    """, (id_equipo,))
    almacenamientos_asignados = cur.fetchall()

    # 🔹 Almacenamiento disponible
    cur.execute("""
        SELECT a.id, a.marca, c.capacidad, c.tipo
        FROM almacenamiento a
        JOIN almacenamientocarac c ON a.idalmacenamientocarac = c.id
        WHERE a.idequipo IS NULL and a.baja = FALSE
    """)
    almacenamientos_disponibles = cur.fetchall()

       # 🔹 memoria ram asignado
    cur.execute("""
        SELECT m.id, m.marca, mr.tipo, mr.capacidad,mr.frecuencia
        FROM memoria m
        JOIN memoriaram mr ON m.idmemoriaram = mr.id
        WHERE m.idequipo = %s
    """, (id_equipo,))
    memoria_asignados = cur.fetchall()

    # 🔹 memoria ram  disponible
    cur.execute("""
        SELECT m.id, m.marca, mr.tipo, mr.capacidad,mr.frecuencia
        FROM memoria m
        JOIN memoriaram mr ON m.idmemoriaram = mr.id
        WHERE m.idequipo IS NULL and m.baja = FALSE
    """)
    memoria_disponibles = cur.fetchall()

       # 🔹 grafica  asignado
    cur.execute("""
        SELECT g.id, g.tipo, g.capacidad, g.marca,g.chipset
        FROM grafica g
        JOIN graficapc gp ON gp.idgrafica  = g.id
        WHERE gp.idequipo = %s
    """, (id_equipo,))
    grafica_asignados = cur.fetchall()

    # 🔹 grafica   disponible
    cur.execute("""
        SELECT g.id, g.tipo, g.capacidad, g.marca,g.chipset
        FROM grafica g
        JOIN graficapc gp ON gp.idgrafica  = g.id
        WHERE gp.idequipo IS NULL and gp.baja = FALSE
    """)
    grafica_disponibles = cur.fetchall()
       # 🔹 mantenimiento  asignado
    cur.execute("""
        SELECT m.id,m.fecharealizado,m.fechaincripcion,m.tipo,m.descripcion,t.cc_nit,t.nombre,t.fecharegistro,string_agg(nt.numero, ', ') AS numero 
        FROM mantenimientoobservacionesugerencias m
        JOIN tecnicos  t ON t.id  = m.idtecnico
        JOIN teletec  nt ON nt.idtecnicos  = m.idtecnico
        WHERE m.idequipo = %s
        GROUP BY 
        m.id, m.fecharealizado, m.fechaincripcion, m.tipo, m.descripcion,
       t.cc_nit, t.nombre, t.fecharegistro       
       """, (id_equipo,))
    mantenimiento_asignados = cur.fetchall()

    # 🔹 mantenimiento   disponible
    cur.execute("""
        SELECT m.id,m.fecharealizado,m.fechaincripcion,m.tipo,m.descripcion,t.cc_nit,t.nombre,t.fecharegistro,string_agg(nt.numero, ', ') AS numero 
        FROM mantenimientoobservacionesugerencias m
        JOIN tecnicos t ON t.id  = m.idtecnico
        JOIN teletec nt ON nt.idtecnicos  = m.idtecnico
        WHERE m.idequipo IS NULL
        GROUP BY 
        m.id, m.fecharealizado, m.fechaincripcion, m.tipo, m.descripcion,
        t.cc_nit, t.nombre, t.fecharegistro  
""")
    mantenimiento_disponibles = cur.fetchall()
           # 🔹 grafica  asignado
    cur.execute("""
        SELECT p.id,mp.marca,p.pulgadas,p.voltaje,p.amperaje,p.serial
        FROM pantalla p
        JOIN marcapantalla mp ON p.idmarcapantalla = mp.id
        WHERE p.idequipo = %s
    """, (id_equipo,))
    pantallas_asignados = cur.fetchall()

    # 🔹 grafica   disponible
    cur.execute("""
        SELECT p.id,mp.marca,p.pulgadas,p.voltaje,p.amperaje,p.serial
        FROM pantalla p
        JOIN marcapantalla mp ON p.idmarcapantalla = mp.id
        WHERE p.idequipo IS NULL and p.baja = FALSE
    """)
    pantallas_disponibles = cur.fetchall()
         # 🔹 perisfericos  asignado
    cur.execute("""
        SELECT p.id,mp.marcas,p.serial,p.modelo,p.descripcion
        FROM perisfericos p
        JOIN marcaperis mp ON p.idmarca = mp.id
        WHERE p.idequipo = %s
    """, (id_equipo,))
    perisfericos_asignados = cur.fetchall()

    # 🔹 perisfericos   disponible
    cur.execute("""
        SELECT p.id,mp.marcas,p.serial,p.modelo,p.descripcion
        FROM perisfericos p
        JOIN marcaperis mp ON p.idmarca = mp.id
        WHERE p.idequipo IS NULL and p.baja = FALSE
    """)
    perisfericos_disponibles = cur.fetchall()
             # 🔹 licencias  asignado
    cur.execute("""
        SELECT id,nombresoftware,fechacompra,fechaexpiracion,clavelicencia
        FROM licenciasequipo  
        WHERE idequipo = %s
    """, (id_equipo,))
    licencias_asignados = cur.fetchall()

    # 🔹 licencias   disponible
    cur.execute("""
        SELECT id,nombresoftware,fechacompra,fechaexpiracion,clavelicencia
        FROM licenciasequipo  
        WHERE idequipo IS NULL AND fechaexpiracion >= CURRENT_DATE
    """)
    licencias_disponibles = cur.fetchall()
    cur.close()
    conn.close()

    return render_template(
        "equiposcaracteristicas.html",
        id_equipo=id_equipo,
        equipo=equipo,
        caracteristicas=caracteristicas,
        boards=boards,
        sistemas=sistemas,
        procesadores=procesadores,
        usuarios_dominio=usuarios_dominio,
        sedes=sedes,
        oficinas=oficinas,
        pisos=pisos,
        usuarios=usuarios,
        almacenamientos_asignados=almacenamientos_asignados,
        almacenamientos_disponibles=almacenamientos_disponibles,
        memoria_asignados=memoria_asignados,
        memoria_disponibles=memoria_disponibles,
        grafica_asignados=grafica_asignados,
        grafica_disponibles=grafica_disponibles,
        mantenimiento_asignados=mantenimiento_asignados,
        mantenimiento_disponibles=mantenimiento_disponibles,
        pantallas_asignados=pantallas_asignados,
        pantallas_disponibles=pantallas_disponibles,
        perisfericos_asignados=perisfericos_asignados,
        perisfericos_disponibles=perisfericos_disponibles,
        licencias_asignados=licencias_asignados,
        licencias_disponibles=licencias_disponibles,    
        usuario=session["usuario"],
        rol=session["rol"],
        rol_id=session["rol_id"]
    )


@caracequipos_bp.route("/equiposcaracteristicas/<int:id_equipo>/editar", methods=["POST"])
def editar_equipo(id_equipo):
    if "usuario" not in session:
        return jsonify({"error": "No autorizado"}), 401

    data = request.form
    foto = request.files.get("foto")  # Aquí viene la foto
    conn = get_connection()
    cur = conn.cursor()

    # Update características
    cur.execute("""
        UPDATE caracteristicasequipo
        SET nombreequipo = %s,
            direccionmac = %s,
            ip = %s,
            observaciones = %s,
            sockets = %s,
            frecuenciaram = %s,
            anidesk = %s,
            idboard = %s,
            idsistemaoperativo = %s,
            idprocesador = %s,
            idusuariodominio = %s
        WHERE idequipo = %s
    """, (
        data.get("nombreequipo") or None,
        data.get("direccionmac") or None,
        data.get("ip") or None,
        data.get("observaciones") or None,
        data.get("sockets") or None,
        data.get("frecuenciaram") or None,
        data.get("anidesk") or None,
        data.get("idboard") or None,
        data.get("idsistemaoperativo") or None,
        data.get("idprocesador") or None,
        data.get("idusuariodominio") or None,
        id_equipo
    ))

# Update equipo (incluyendo foto si la mandaron)
    if foto and foto.filename:
        
        cur.execute("""
            UPDATE equipo
            SET estado = %s,
                idsede = %s,
                idoficina = %s,
                idpiso = %s,
                idusuario = %s,
                foto = %s
            WHERE id = %s
        """, (
            data.get("estado") or None,
            data.get("idsede") or None,
            data.get("idoficina") or None,
            data.get("idpiso") or None,
            data.get("idusuario") or None,
            foto.read(),   # Guarda como BYTEA
            id_equipo
        ))
    else:
        cur.execute("""
        UPDATE equipo
        SET estado = %s,
            idsede = %s,
            idoficina = %s,
            idpiso = %s,
            idusuario = %s
           WHERE id = %s
          """, (
          data.get("estado") or None,
          data.get("idsede") or None,
         data.get("idoficina") or None,
          data.get("idpiso") or None,
           data.get("idusuario") or None,
         id_equipo
          ))

    conn.commit()
    cur.close()
    conn.close()

    flash("Equipo actualizado correctamente", "success")
    return redirect(url_for("caracequipos.equiposcaracteristicas", id_equipo=id_equipo))

# 🔹 Acción dar de baja
@caracequipos_bp.route("/equiposcaracteristicas/<int:id_equipo>/baja", methods=["POST"])
def baja_equipo(id_equipo):
    if "usuario" not in session:
        return jsonify({"error": "No autorizado"}), 401

    descripcion = request.form.get("descripcion")
    fecha_baja = request.form.get("fecha_baja") or date.today()

    try:
        conn = get_connection()
        cur = conn.cursor()
        cur.execute("CALL sp_baja_equipo(%s, %s, %s)", (id_equipo, descripcion, fecha_baja))
        conn.commit()
        cur.close()
        conn.close()

        flash("Equipo dado de baja exitosamente", "success")
    except Exception as e:
        flash(f"No se pudo dar de baja: {str(e)}", "danger")

    return redirect(url_for("caracequipos.equiposcaracteristicas", id_equipo=id_equipo))

@caracequipos_bp.route("/equiposcaracteristicas/<int:id_equipo>/almacenamiento/agregar", methods=["POST"])
def agregar_almacenamiento(id_equipo):
    if "usuario" not in session:
        return jsonify({"error": "No autorizado"}), 401

    idalmacenamiento = request.form.get("idalmacenamiento")
    conn = get_connection()
    cur = conn.cursor()
    cur.execute("UPDATE almacenamiento SET idequipo = %s WHERE id = %s", (id_equipo, idalmacenamiento))
    conn.commit()
    cur.close()
    conn.close()

    flash("Almacenamiento asignado correctamente", "success")
    return redirect(url_for("caracequipos.equiposcaracteristicas", id_equipo=id_equipo))

# 🔹 Quitar almacenamiento
@caracequipos_bp.route("/equiposcaracteristicas/<int:id_equipo>/almacenamiento/<int:idalmacenamiento>/quitar", methods=["POST"])
def quitar_almacenamiento(id_equipo, idalmacenamiento):
    if "usuario" not in session:
        return jsonify({"error": "No autorizado"}), 401

    conn = get_connection()
    cur = conn.cursor()
    cur.execute("UPDATE almacenamiento SET idequipo = NULL WHERE id = %s", (idalmacenamiento,))
    conn.commit()
    cur.close()
    conn.close()

    flash("Almacenamiento quitado correctamente", "info")
    return redirect(url_for("caracequipos.equiposcaracteristicas", id_equipo=id_equipo))

# 🔹 Agregar memoria ram
@caracequipos_bp.route("/equiposcaracteristicas/<int:id_equipo>/memoria/agregar", methods=["POST"])
def agregar_memoria(id_equipo):
    if "usuario" not in session:
        return jsonify({"error": "No autorizado"}), 401

    idmemoriaram = request.form.get("idmemoriaram")
    conn = get_connection()
    cur = conn.cursor()
   # 1️⃣ Ver cuántos sockets tiene el equipo
    cur.execute("SELECT sockets FROM caracteristicasequipo WHERE idequipo = %s", (id_equipo,))
    result = cur.fetchone()
    if not result:
        flash("No se encontraron características del equipo", "danger")
        cur.close()
        conn.close()
        return redirect(url_for("caracequipos.equiposcaracteristicas", id_equipo=id_equipo))

    max_sockets = result[0]

    # 2️⃣ Ver cuántas memorias ya tiene asignadas
    cur.execute("SELECT COUNT(*) FROM memoria WHERE idequipo = %s", (id_equipo,))
    usadas = cur.fetchone()[0]

    if usadas >= max_sockets:
        flash(f"El equipo ya tiene el máximo de {max_sockets} memorias asignadas", "warning")
        cur.close()
        conn.close()
        return redirect(url_for("caracequipos.equiposcaracteristicas", id_equipo=id_equipo))

    # 3️⃣ Si todavía hay sockets disponibles → asignar
    cur.execute("UPDATE memoria SET idequipo = %s WHERE id = %s", (id_equipo, idmemoriaram))
    conn.commit()
    cur.close()
    conn.close()

    flash("Memoria RAM asignada correctamente", "success")
    return redirect(url_for("caracequipos.equiposcaracteristicas", id_equipo=id_equipo))

# 🔹 Quitar almacenamiento
@caracequipos_bp.route("/equiposcaracteristicas/<int:id_equipo>/memoria/<int:idmemoriaram>/quitar", methods=["POST"])
def quitar_memoria(id_equipo, idmemoriaram):
    if "usuario" not in session:
        return jsonify({"error": "No autorizado"}), 401

    conn = get_connection()
    cur = conn.cursor()
    cur.execute("UPDATE memoria SET idequipo = NULL WHERE id = %s", (idmemoriaram,))
    conn.commit()
    cur.close()
    conn.close()

    flash("Memoria Ram quitado correctamente", "info")
    return redirect(url_for("caracequipos.equiposcaracteristicas", id_equipo=id_equipo))
# 🔹 Agregar grafica
@caracequipos_bp.route("/equiposcaracteristicas/<int:id_equipo>/grafica/agregar", methods=["POST"])
def agregar_grafica(id_equipo):
    if "usuario" not in session:
        return jsonify({"error": "No autorizado"}), 401

    idgrafica = request.form.get("idgrafica")
    conn = get_connection()
    cur = conn.cursor()
    cur.execute("UPDATE graficapc SET idequipo = %s WHERE idgrafica = %s", (id_equipo, idgrafica))
    conn.commit()
    cur.close()
    conn.close()

    flash("grafica asignado correctamente", "success")
    return redirect(url_for("caracequipos.equiposcaracteristicas", id_equipo=id_equipo))

# 🔹 Quitar grafica
@caracequipos_bp.route("/equiposcaracteristicas/<int:id_equipo>/grafica/<int:idgrafica>/quitar", methods=["POST"])
def quitar_grafica(id_equipo, idgrafica):
    if "usuario" not in session:
        return jsonify({"error": "No autorizado"}), 401

    conn = get_connection()
    cur = conn.cursor()
    cur.execute("UPDATE graficapc SET idequipo = NULL WHERE idgrafica = %s", (idgrafica,))
    conn.commit()
    cur.close()
    conn.close()

    flash("grafica quitado correctamente", "info")
    return redirect(url_for("caracequipos.equiposcaracteristicas", id_equipo=id_equipo))
# 🔹 Agregar mantenimiento
@caracequipos_bp.route("/equiposcaracteristicas/<int:id_equipo>/mantenimiento/agregar", methods=["POST"])
def agregar_mantenimiento(id_equipo):
    if "usuario" not in session:
        return jsonify({"error": "No autorizado"}), 401

    idmantenimiento = request.form.get("idmantenimiento")
    conn = get_connection()
    cur = conn.cursor()
    cur.execute("UPDATE mantenimientoobservacionesugerencias  SET idequipo = %s WHERE id = %s", (id_equipo, idmantenimiento))
    conn.commit()
    cur.close()
    conn.close()

    flash("mantenimiento asignado correctamente", "success")
    return redirect(url_for("caracequipos.equiposcaracteristicas", id_equipo=id_equipo))

# 🔹 Quitar mantenimiento
@caracequipos_bp.route("/equiposcaracteristicas/<int:id_equipo>/mantenimiento/<int:idmantenimiento>/quitar", methods=["POST"])
def quitar_mantenimiento(id_equipo, idmantenimiento):
    if "usuario" not in session:
        return jsonify({"error": "No autorizado"}), 401

    conn = get_connection()
    cur = conn.cursor()
    cur.execute("UPDATE mantenimientoobservacionesugerencias  SET idequipo = NULL WHERE id = %s", (idmantenimiento,))
    conn.commit()
    cur.close()
    conn.close()

    flash("mantenimiento quitado correctamente", "info")
    return redirect(url_for("caracequipos.equiposcaracteristicas", id_equipo=id_equipo))
# 🔹 Agregar pantallas
@caracequipos_bp.route("/equiposcaracteristicas/<int:id_equipo>/pantallas/agregar", methods=["POST"])
def agregar_pantallas(id_equipo):
    if "usuario" not in session:
        return jsonify({"error": "No autorizado"}), 401

    idpantallas = request.form.get("idpantallas")
    conn = get_connection()
    cur = conn.cursor()
    cur.execute("UPDATE pantalla  SET idequipo = %s WHERE id = %s", (id_equipo, idpantallas))
    conn.commit()
    cur.close()
    conn.close()

    flash("pantallas asignado correctamente", "success")
    return redirect(url_for("caracequipos.equiposcaracteristicas", id_equipo=id_equipo))

# 🔹 Quitar  pantallas

@caracequipos_bp.route("/equiposcaracteristicas/<int:id_equipo>/pantallas/<int:idpantallas>/quitar", methods=["POST"])
def quitar_pantallas(id_equipo, idpantallas):
    if "usuario" not in session:
        return jsonify({"error": "No autorizado"}), 401

    conn = get_connection()
    cur = conn.cursor()
    cur.execute("UPDATE pantalla  SET idequipo = NULL WHERE id = %s", (idpantallas,))
    conn.commit()
    cur.close()
    conn.close()

    flash("pantallas quitado correctamente", "info")
    return redirect(url_for("caracequipos.equiposcaracteristicas", id_equipo=id_equipo))
# 🔹 Agregar perisfericos
@caracequipos_bp.route("/equiposcaracteristicas/<int:id_equipo>/perisfericos/agregar", methods=["POST"])
def agregar_perisfericos(id_equipo):
    if "usuario" not in session:
        return jsonify({"error": "No autorizado"}), 401

    idperisfericos = request.form.get("idperisfericos")
    conn = get_connection()
    cur = conn.cursor()
    cur.execute("UPDATE perisfericos  SET idequipo = %s WHERE id = %s", (id_equipo, idperisfericos))
    conn.commit()
    cur.close()
    conn.close()

    flash("perisfericos asignado correctamente", "success")
    return redirect(url_for("caracequipos.equiposcaracteristicas", id_equipo=id_equipo))

# 🔹 Quitar  perisfericos

@caracequipos_bp.route("/equiposcaracteristicas/<int:id_equipo>/perisfericos/<int:idperisfericos>/quitar", methods=["POST"])
def quitar_perisfericos(id_equipo, idperisfericos):
    if "usuario" not in session:
        return jsonify({"error": "No autorizado"}), 401

    conn = get_connection()
    cur = conn.cursor()
    cur.execute("UPDATE perisfericos  SET idequipo = NULL WHERE id= %s", (idperisfericos,))
    conn.commit()
    cur.close()
    conn.close()

    flash("perisfericos quitado correctamente", "info")
    return redirect(url_for("caracequipos.equiposcaracteristicas", id_equipo=id_equipo))
# 🔹 Agregar perisfericos
@caracequipos_bp.route("/equiposcaracteristicas/<int:id_equipo>/licencias/agregar", methods=["POST"])
def agregar_licencias(id_equipo):
    if "usuario" not in session:
        return jsonify({"error": "No autorizado"}), 401

    idlicencias = request.form.get("idlicencias")
    conn = get_connection()
    cur = conn.cursor()
    cur.execute("UPDATE licenciasequipo   SET idequipo = %s WHERE id = %s", (id_equipo, idlicencias))
    conn.commit()
    cur.close()
    conn.close()

    flash("licencias asignado correctamente", "success")
    return redirect(url_for("caracequipos.equiposcaracteristicas", id_equipo=id_equipo))

# 🔹 Quitar  perisfericos

@caracequipos_bp.route("/equiposcaracteristicas/<int:id_equipo>/licencias/<int:idlicencias>/quitar", methods=["POST"])
def quitar_licencias(id_equipo, idlicencias):
    if "usuario" not in session:
        return jsonify({"error": "No autorizado"}), 401

    conn = get_connection()
    cur = conn.cursor()
    cur.execute("UPDATE licenciasequipo   SET idequipo = NULL WHERE id= %s", (idlicencias,))
    conn.commit()
    cur.close()
    conn.close()

    flash("licencias quitado correctamente", "info")
    return redirect(url_for("caracequipos.equiposcaracteristicas", id_equipo=id_equipo))

# ===========================
# 🔹 Exportar PDF
# ===========================
@caracequipos_bp.route("/equiposcaracteristicas/<int:id_equipo>/export/pdf")
def export_pdf(id_equipo):
    from datetime import datetime
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_RIGHT
    conn = get_connection()
    cur = conn.cursor(cursor_factory=extras.RealDictCursor)

    # Consultas
    cur.execute("""SELECT e.codigo,e.tipo,e.marca,e.modelo,e.serial,
                          e.fechacompra,e.estado,e.baja,e.fecharegistro,
                          s.sede, o.oficina_salon, p.piso, u.nombre as usuario_nombre,u.cc,u.email,ar.area,ca.descripcion as cargo
                   FROM equipo e
                   LEFT JOIN sede s ON e.idsede = s.id
                   LEFT JOIN oficina o ON e.idoficina = o.id
                   LEFT JOIN piso p ON e.idpiso = p.id
                   LEFT JOIN usuario u ON e.idusuario = u.id
                   LEFT JOIN cargo ca ON ca.id = u.idcargo 
                   LEFT JOIN area ar ON ar.id = u.idarea
                   WHERE e.id = %s""", (id_equipo,))
    equipo = cur.fetchone()
    cur.execute ("select foto from equipo where id = %s", (id_equipo,))
    foto = cur.fetchone()

    cur.execute("""SELECT c.id,c.direccionmac,c.ip,c.observaciones,c.nombreequipo,c.sockets,c.frecuenciaram,c.anidesk,
                 b.tarjetamadre,b.chipset, so.descripcion as sistema,so.arquitectura, pr.modelo as procesador,pr.marca,pr.frecuencia,pr.nucleos,pr.plataforma,
               ud.nombre as usuario_dominio
        FROM caracteristicasequipo c
        LEFT JOIN board b ON c.idboard = b.id
        LEFT JOIN sistemaopera so ON c.idsistemaoperativo = so.id
        LEFT JOIN procesador pr ON c.idprocesador = pr.id
        LEFT JOIN usuariodominio ud ON c.idusuariodominio = ud.id
        WHERE c.idequipo  = %s""", (id_equipo,))
    caracteristicas = cur.fetchone()

    cur.execute("""SELECT marca, horas_uso, estado, capacidad, tipo
                   FROM almacenamiento a
                   JOIN almacenamientocarac c ON a.idalmacenamientocarac = c.id
                   WHERE a.idequipo = %s""", (id_equipo,))
    almacenamientos = cur.fetchall()

    cur.execute("""SELECT marca, tipo, capacidad, frecuencia
                   FROM memoria m
                   JOIN memoriaram mr ON m.idmemoriaram = mr.id
                   WHERE m.idequipo = %s""", (id_equipo,))
    memorias = cur.fetchall()
    cur.execute("""SELECT g.tipo, g.capacidad, g.marca, g.chipset
                   FROM grafica g
                   JOIN graficapc gr ON gr.idgrafica  = g.id
                   WHERE gr.idequipo = %s""", (id_equipo,))
    grafica = cur.fetchall()
    cur.execute("""
        SELECT m.fecharealizado,m.fechaincripcion,m.tipo,m.descripcion,t.cc_nit,t.nombre,t.fecharegistro,string_agg(nt.numero, ', ') AS telefonos
        FROM mantenimientoobservacionesugerencias  m
        JOIN tecnicos  t ON t.id  = m.idtecnico
        JOIN teletec  nt ON nt.idtecnicos  = m.idtecnico
        WHERE m.idequipo = %s
        GROUP BY 
        m.fecharealizado,
        m.fechaincripcion,
        m.tipo,
        m.descripcion,
        t.cc_nit,
        t.nombre,
        t.fecharegistro
    """, (id_equipo,))
    mantenimiento = cur.fetchall()
    cur.execute("""
        SELECT p.id,mp.marcas,p.serial,p.modelo,p.descripcion
        FROM perisfericos p
        JOIN marcaperis mp ON p.idmarca = mp.id
        WHERE p.idequipo = %s
    """, (id_equipo,))
    perisfericos = cur.fetchall()
    cur.execute("""
        SELECT mp.marca,p.pulgadas,p.voltaje,p.amperaje,p.serial
        FROM pantalla p
        JOIN marcapantalla mp ON p.idmarcapantalla = mp.id
        WHERE p.idequipo = %s
    """, (id_equipo,))
    pantallas = cur.fetchall()
    cur.execute("""
        SELECT id,nombresoftware,fechacompra,fechaexpiracion,clavelicencia
        FROM licenciasequipo  
        WHERE idequipo = %s
    """, (id_equipo,))
    licencias= cur.fetchall()
    baja_info = None
    if equipo and equipo["baja"]:
        cur = conn.cursor(cursor_factory=extras.RealDictCursor)
        cur.execute("""SELECT id, tabla, idregistro, descripcion, fecharegistro, fechabaja 
                        FROM baja 
                        WHERE idregistro = %s AND tabla = 'equipo'""", (id_equipo,))
        baja_info = cur.fetchone()

    cur.close()
    conn.close()

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                            leftMargin=40, rightMargin=40,
                            topMargin=60, bottomMargin=40)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="SectionTitle", fontSize=14, textColor="#2C3E50", spaceAfter=10, spaceBefore=10, fontName="Helvetica-Bold"))
    styles.add(ParagraphStyle(name="Footer", fontSize=8, alignment=TA_RIGHT, textColor=colors.grey))
    styles["Title"].fontSize = 18
    styles["Title"].textColor = "#4A90E2"
    elements = []
    def separador():
        return Table([[""]], colWidths=[500], style=TableStyle([
            ("LINEBELOW", (0, 0), (-1, -1), 0.25, colors.grey)
        
        ]))
    def safe_value(v):
        if v is None or str(v).strip() == "":
             return "No asignado"
        if isinstance(v, datetime):
             return v.strftime("%d/%m/%Y")
        return str(v)


    # === Encabezado con logo ===
    logo_path = "static/LOGO.png"
    header_data = []
    if os.path.exists(logo_path):
        header_data.append([Image(logo_path, width=40, height=40),
                            Paragraph("<b>Reporte Técnico de Equipo</b>", styles["Title"])])
        header_table = Table(header_data, colWidths=[50, 400])
        header_table.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "MIDDLE")]))
        elements.append(header_table)
        elements.append(Spacer(1, 20))

    # === Foto del equipo ===
    if foto and foto["foto"]:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp_img:
            tmp_img.write(foto["foto"])
            tmp_img.flush()
            escala = 0.5
            foto_equipo = Image(tmp_img.name, width=8*cm * escala, height=6*cm * escala)
            foto_equipo.hAlign = "CENTER"
            bordered = Table([[foto_equipo]], style=TableStyle([
                ("BOX", (0, 0), (-1, -1), 1, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE")
            ]))
            elements.append(bordered)
            elements.append(Spacer(1, 10))

    if equipo and equipo["baja"] and baja_info:
         elements.append(Paragraph("📉 Información de Baja", styles["SectionTitle"]))
         elements.append(separador())
         data = [["Campo", "Valor"]]
         for k, v in baja_info.items():
             if isinstance(v, datetime):
                 v = v.strftime("%d/%m/%Y")
             data.append([str(k).capitalize(), str(v)])
         table = Table(data, colWidths=[150, 320])
         table.setStyle(TableStyle([
             ("BACKGROUND", (0, 0), (-1, 0), "#E74C3C"),
             ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
             ("BOX", (0, 0), (-1, -1), 0.75, colors.grey),
             ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
         ]))
         elements.append(table)
         elements.append(Spacer(1, 20))


    # === Información General ===
    elements.append(Paragraph("🖥️ Información General", styles["SectionTitle"]))
    elements.append(separador())
    data = [["Campo", "Valor"]]
    for k, v in equipo.items():
       data.append([str(k).capitalize(), safe_value(v)]) 
    table = Table(data, colWidths=[150, 320])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), "#4A90E2"),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("BOX", (0, 0), (-1, -1), 0.75, colors.grey),
        ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
    ]))
    elements.append(table)
    elements.append(Spacer(1, 20))


    # === Características ===
    elements.append(Paragraph("⚙️ Características Técnicas", styles["SectionTitle"]))
    elements.append(separador())
    data = [["Campo", "Valor"]]
    for k, v in caracteristicas.items():
        data.append([str(k).capitalize(), safe_value(v)])
    table = Table(data, colWidths=[150, 320])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), "#4A90E2"),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("BOX", (0, 0), (-1, -1), 0.75, colors.grey),
        ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
    ]))
    elements.append(table)
    elements.append(Spacer(1, 20))

    # === Almacenamiento ==================================================================================================================
    elements.append(Paragraph("💾 Almacenamiento", styles["SectionTitle"]))
    elements.append(separador())

    if almacenamientos:
        data = [["Marca", "Capacidad", "Tipo", "Horas Uso", "Estado"]]
        estado_iconos = {
        "bueno": "🟢 Bueno",
        "regular": "🟠 Regular",
        "malo": "🔴 Malo"
        }
        for a in almacenamientos:
            estado_raw = a["estado"].lower() if a["estado"] else None
            estado_visual = estado_iconos.get(estado_raw, "No asignado")
            data.append([
                safe_value(a["marca"]),
                f"{a['capacidad']} GB" if a["capacidad"] else "No asignado",
                safe_value(a["tipo"]),
                f"{a['horas_uso']} Horas" if a["horas_uso"] else "No asignado",
                estado_visual
            ])
        table = Table(data, colWidths=[80, 80, 80, 80, 100])
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), "#4A90E2"),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("BOX", (0, 0), (-1, -1), 0.75, colors.grey),
            ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ]))
        elements.append(table)
    else:
        elements.append(Paragraph("⚠️ No hay almacenamiento asignado", styles["Normal"]))
    elements.append(Spacer(1, 20))


# === Memoria RAM =====================================================================================================================
    elements.append(Paragraph("🧠 Memoria RAM", styles["SectionTitle"]))
    elements.append(separador())

    if memorias:
        data = [["Marca", "Tipo", "Capacidad", "Frecuencia"]]
        for m in memorias:
            data.append([
                safe_value(m["marca"]),
                safe_value(m["tipo"]),
                f"{m['capacidad']} GB" if m["capacidad"] else "No asignado",
                f"{m['frecuencia']} MHz" if m["frecuencia"] else "No asignado"
            ])
        table = Table(data, colWidths=[100, 100, 100, 100])
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), "#4A90E2"),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("BOX", (0, 0), (-1, -1), 0.75, colors.grey),
            ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ]))
        elements.append(table)
    else:
       elements.append(Paragraph("⚠️ No hay memoria RAM asignada", styles["Normal"]))
    elements.append(Spacer(1, 20))
# === grafica ==============================================================================================================================

    elements.append(Paragraph(" Grafica", styles["SectionTitle"]))
    elements.append(separador())

    if grafica:
        data = [["Tipo", "Capacidad", "Marca", "Chipset "]]
        for g in grafica:
            data.append([
                safe_value(g["tipo"]),
                f"{g['capacidad']} GB" if g["capacidad"] else "No asignado",
                safe_value(g["marca"]),
                safe_value(g["chipset"]),
            ])
        table = Table(data, colWidths=[100, 100, 100, 100])
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), "#4A90E2"),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("BOX", (0, 0), (-1, -1), 0.75, colors.grey),
            ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ]))
        elements.append(table)
    else:
       elements.append(Paragraph("⚠️ No hay grafica asignada", styles["Normal"]))
    elements.append(Spacer(1, 20))
# === mantenimiento ==============================================================================================================================

    elements.append(Paragraph("🛠️ Mantenimiento", styles["SectionTitle"]))
    elements.append(separador())

    if mantenimiento:
        data_mant  = [["Fecha realizado", "Fecha inscripción", "Tipo", "Descripción"]]
        for m in mantenimiento:
            data_mant.append([
                safe_value(m["fecharealizado"]),
                safe_value(m["fechaincripcion"]),
                safe_value(m["tipo"]),
                safe_value(m["descripcion"])
                
            ])
        table_mant = Table(data_mant, colWidths=[80, 100, 80, 180])
        table_mant.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), "#4A90E2"),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("BOX", (0, 0), (-1, -1), 0.75, colors.grey),
            ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
        ]))
        elements.append(table_mant)
        elements.append(Spacer(1, 10))
        data_tec = [["CC/NIT", "Nombre técnico", "Fecha registro técnico", "Teléfonos"]]
        for m in mantenimiento:
            telefonos = safe_value(m["telefonos"]) if m.get("telefonos") else "No asignado"
            data_tec.append([
                safe_value(m["cc_nit"]),
                safe_value(m["nombre"]),
                safe_value(m["fecharegistro"]),
                telefonos
                
            ])
        table_tec = Table(data_tec, colWidths=[80, 120, 110, 140])
        table_tec.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), "#4A90E2"),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("BOX", (0, 0), (-1, -1), 0.75, colors.grey),
            ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
        ]))
        elements.append(table_tec)
        
    else:
       elements.append(Paragraph("⚠️ No hay mantenimiento asignada", styles["Normal"]))
    elements.append(Spacer(1, 20))
# === pantallas ==============================================================================================================================
    elements.append(Paragraph("pantallas", styles["SectionTitle"]))
    elements.append(separador())

    if pantallas:
        datap =[["marca", "pulgadas", "voltaje", "amperaje ","serial"]]
        for p in pantallas:
            datap.append([
                safe_value(p["marca"]),
                safe_value(p["pulgadas"]),
                safe_value(p["voltaje"]),
                safe_value(p["amperaje"]),
                safe_value(p["serial"]),
            ])
        tablep = Table(datap, colWidths=[100, 100, 100, 100])
        tablep.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), "#4A90E2"),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("BOX", (0, 0), (-1, -1), 0.75, colors.grey),
            ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ]))
        elements.append(tablep)
    else:
       elements.append(Paragraph("⚠️ No hay pantallas asignada", styles["Normal"]))
    elements.append(Spacer(1, 20))

# === perisfericos ==============================================================================================================================
    elements.append(Paragraph("perisfericos", styles["SectionTitle"]))
    elements.append(separador())

    if perisfericos:
        datapr =[["marca", "serial", "modelo", "descripcion"]]
        for pr in perisfericos:
            datapr.append([
                safe_value(pr["marcas"]),
                safe_value(pr["serial"]),
                safe_value(pr["modelo"]),
                safe_value(pr["descripcion"]),
            ])
        tablepr = Table(datapr, colWidths=[100, 100, 100, 100])
        tablepr.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), "#4A90E2"),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("BOX", (0, 0), (-1, -1), 0.75, colors.grey),
            ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ]))
        elements.append(tablepr)
    else:
       elements.append(Paragraph("⚠️ No hay perisfericos asignada", styles["Normal"]))
    elements.append(Spacer(1, 20))
# === licencias ==============================================================================================================================
    elements.append(Paragraph("licencias", styles["SectionTitle"]))
    elements.append(separador())

    if licencias:
        datal =[["nombre", "fecha compra ", "fecha expiracion ", "clave licencia"]]
        for l in licencias:
            datal.append([
                safe_value(l["nombresoftware"]),
                safe_value(l["fechacompra"]),
                safe_value(l["fechaexpiracion"]),
                safe_value(l["clavelicencia"]),
            ])
        tablel = Table(datal, colWidths=[100, 100, 100, 100])
        tablel.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), "#4A90E2"),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("BOX", (0, 0), (-1, -1), 0.75, colors.grey),
            ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ]))
        elements.append(tablel)
    else:
       elements.append(Paragraph("⚠️ No hay licencias asignada", styles["Normal"]))
    elements.append(Spacer(1, 20))
# === Construir el PDF ===
    doc.build(elements)
    buffer.seek(0)

    # === Retornar respuesta ===
    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"equipo_{id_equipo}.pdf",
        mimetype="application/pdf"
    )

# ===========================
# 🔹 Exportar Word
# ===========================
@caracequipos_bp.route("/equiposcaracteristicas/<int:id_equipo>/export/word")
def export_word(id_equipo):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO
    from datetime import datetime

    conn = get_connection()
    cur = conn.cursor(cursor_factory=extras.RealDictCursor)

    # === Consultas principales ===
    cur.execute("""SELECT e.codigo,e.tipo,e.marca,e.modelo,e.serial,
                          e.fechacompra,e.estado,e.baja,e.fecharegistro,
                          s.sede, o.oficina_salon, p.piso, u.nombre as usuario_nombre,u.cc,u.email,ar.area,ca.descripcion as cargo
                   FROM equipo e
                   LEFT JOIN sede s ON e.idsede = s.id
                   LEFT JOIN oficina o ON e.idoficina = o.id
                   LEFT JOIN piso p ON e.idpiso = p.id
                   LEFT JOIN usuario u ON e.idusuario = u.id
                   LEFT JOIN cargo ca ON ca.id = u.idcargo 
                   LEFT JOIN area ar ON ar.id = u.idarea
                   WHERE e.id = %s""", (id_equipo,))
    equipo = cur.fetchone()

    cur.execute("""SELECT c.id,c.direccionmac,c.ip,c.observaciones,c.nombreequipo,c.sockets,c.frecuenciaram,c.anidesk,
                 b.tarjetamadre,b.chipset, so.descripcion as sistema,so.arquitectura, pr.modelo as procesador,pr.marca,pr.frecuencia,pr.nucleos,pr.plataforma,
               ud.nombre as usuario_dominio
        FROM caracteristicasequipo c
        LEFT JOIN board b ON c.idboard = b.id
        LEFT JOIN sistemaopera so ON c.idsistemaoperativo = so.id
        LEFT JOIN procesador pr ON c.idprocesador = pr.id
        LEFT JOIN usuariodominio ud ON c.idusuariodominio = ud.id
        WHERE c.idequipo  = %s""", (id_equipo,))
    caracteristicas = cur.fetchone()

    cur.execute("""SELECT marca, horas_uso, estado, capacidad, tipo
                   FROM almacenamiento a
                   JOIN almacenamientocarac c ON a.idalmacenamientocarac = c.id
                   WHERE a.idequipo = %s""", (id_equipo,))
    almacenamientos = cur.fetchall()

    cur.execute("""SELECT marca, tipo, capacidad, frecuencia
                   FROM memoria m
                   JOIN memoriaram mr ON m.idmemoriaram = mr.id
                   WHERE m.idequipo = %s""", (id_equipo,))
    memorias = cur.fetchall()

    cur.execute("""SELECT g.tipo, g.capacidad, g.marca, g.chipset
                   FROM grafica g
                   JOIN graficapc gr ON gr.idgrafica  = g.id
                   WHERE gr.idequipo = %s""", (id_equipo,))
    graficas = cur.fetchall()

    cur.execute("""SELECT m.fecharealizado,m.fechaincripcion,m.tipo,m.descripcion,t.cc_nit,t.nombre,t.fecharegistro,string_agg(nt.numero, ', ') AS telefonos
        FROM mantenimientoobservacionesugerencias  m
        JOIN tecnicos  t ON t.id  = m.idtecnico
        JOIN teletec  nt ON nt.idtecnicos  = m.idtecnico
        WHERE m.idequipo = %s
        GROUP BY 
        m.fecharealizado,
        m.fechaincripcion,
        m.tipo,
        m.descripcion,
        t.cc_nit,
        t.nombre,
        t.fecharegistro
    """, (id_equipo,))
    mantenimiento = cur.fetchall()

    cur.execute("""SELECT p.id,mp.marcas,p.serial,p.modelo,p.descripcion
        FROM perisfericos p
        JOIN marcaperis mp ON p.idmarca = mp.id
        WHERE p.idequipo = %s""", (id_equipo,))
    perisfericos = cur.fetchall()

    cur.execute("""SELECT mp.marca,p.pulgadas,p.voltaje,p.amperaje,p.serial
        FROM pantalla p
        JOIN marcapantalla mp ON p.idmarcapantalla = mp.id
        WHERE p.idequipo = %s""", (id_equipo,))
    pantallas = cur.fetchall()

    cur.execute("""SELECT id,nombresoftware,fechacompra,fechaexpiracion,clavelicencia
        FROM licenciasequipo  
        WHERE idequipo = %s""", (id_equipo,))
    licencias = cur.fetchall()

    cur.close()
    conn.close()

    # === Crear documento Word ===
    doc = Document()

    def add_section(title, data_dict_or_list, headers=None):
        # Título de sección
        p = doc.add_paragraph(title)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(42, 90, 170)

        doc.add_paragraph("")  # Espacio

        if isinstance(data_dict_or_list, dict) and data_dict_or_list:
            table = doc.add_table(rows=1, cols=2)
            table.style = "Light List Accent 1"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text, hdr_cells[1].text = "Campo", "Valor"
            for k, v in data_dict_or_list.items():
                val = v.strftime("%d/%m/%Y") if isinstance(v, datetime) else (v if v else "No asignado")
                row_cells = table.add_row().cells
                row_cells[0].text = str(k)
                row_cells[1].text = str(val)
        elif isinstance(data_dict_or_list, list) and data_dict_or_list:
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Light List Accent 1"
            hdr_cells = table.rows[0].cells
            for i, h in enumerate(headers):
                hdr_cells[i].text = h
            for row in data_dict_or_list:
                row_cells = table.add_row().cells
                for i, h in enumerate(headers):
                    val = row.get(h, "No asignado")
                    if isinstance(val, datetime):
                        val = val.strftime("%d/%m/%Y")
                    row_cells[i].text = str(val)
        else:
            doc.add_paragraph("⚠️ No hay información registrada")

        doc.add_paragraph("")  # Espacio

    # === Secciones ===
    add_section("🖥️ Información General", equipo)
    add_section("⚙️ Características Técnicas", caracteristicas)
    add_section("💾 Almacenamiento", almacenamientos, ["marca", "capacidad", "tipo", "horas_uso", "estado"])
    add_section("🧠 Memoria RAM", memorias, ["marca", "tipo", "capacidad", "frecuencia"])
    add_section("🎮 Gráficas", graficas, ["tipo", "capacidad", "marca", "chipset"])
    add_section("🛠️ Mantenimiento", mantenimiento, ["fecharealizado", "fechaincripcion", "tipo", "descripcion", "cc_nit", "nombre", "fecharegistro", "telefonos"])
    add_section("⌨️ Periféricos", perisfericos, ["marcas", "serial", "modelo", "descripcion"])
    add_section("🖥️ Pantallas", pantallas, ["marca", "pulgadas", "voltaje", "amperaje", "serial"])
    add_section("📄 Licencias", licencias, ["nombresoftware", "fechacompra", "fechaexpiracion", "clavelicencia"])

    # === Guardar en buffer ===
    output = BytesIO()
    doc.save(output)
    output.seek(0)

    return send_file(
        output,
        as_attachment=True,
        download_name=f"equipo_{id_equipo}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# ===========================
# 🔹 Exportar Excel
# ===========================
@caracequipos_bp.route("/equiposcaracteristicas/<int:id_equipo>/export/excel")
def export_excel(id_equipo):
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill
    from io import BytesIO
    from datetime import datetime

    conn = get_connection()
    cur = conn.cursor(cursor_factory=extras.RealDictCursor)

    # === Consultas principales ===
    cur.execute("""SELECT e.codigo,e.tipo,e.marca,e.modelo,e.serial,
                          e.fechacompra,e.estado,e.baja,e.fecharegistro,
                          s.sede, o.oficina_salon, p.piso, u.nombre as usuario_nombre,u.cc,u.email,ar.area,ca.descripcion as cargo
                   FROM equipo e
                   LEFT JOIN sede s ON e.idsede = s.id
                   LEFT JOIN oficina o ON e.idoficina = o.id
                   LEFT JOIN piso p ON e.idpiso = p.id
                   LEFT JOIN usuario u ON e.idusuario = u.id
                   LEFT JOIN cargo ca ON ca.id = u.idcargo 
                   LEFT JOIN area ar ON ar.id = u.idarea
                   WHERE e.id = %s""", (id_equipo,))
    equipo = cur.fetchone()

    cur.execute("""SELECT c.id,c.direccionmac,c.ip,c.observaciones,c.nombreequipo,c.sockets,c.frecuenciaram,c.anidesk,
                 b.tarjetamadre,b.chipset, so.descripcion as sistema,so.arquitectura, pr.modelo as procesador,pr.marca,pr.frecuencia,pr.nucleos,pr.plataforma,
               ud.nombre as usuario_dominio
        FROM caracteristicasequipo c
        LEFT JOIN board b ON c.idboard = b.id
        LEFT JOIN sistemaopera so ON c.idsistemaoperativo = so.id
        LEFT JOIN procesador pr ON c.idprocesador = pr.id
        LEFT JOIN usuariodominio ud ON c.idusuariodominio = ud.id
        WHERE c.idequipo  = %s""", (id_equipo,))
    caracteristicas = cur.fetchone()

    cur.execute("""SELECT marca, horas_uso, estado, capacidad, tipo
                   FROM almacenamiento a
                   JOIN almacenamientocarac c ON a.idalmacenamientocarac = c.id
                   WHERE a.idequipo = %s""", (id_equipo,))
    almacenamientos = cur.fetchall()

    cur.execute("""SELECT marca, tipo, capacidad, frecuencia
                   FROM memoria m
                   JOIN memoriaram mr ON m.idmemoriaram = mr.id
                   WHERE m.idequipo = %s""", (id_equipo,))
    memorias = cur.fetchall()

    cur.execute("""SELECT g.tipo, g.capacidad, g.marca, g.chipset
                   FROM grafica g
                   JOIN graficapc gr ON gr.idgrafica  = g.id
                   WHERE gr.idequipo = %s""", (id_equipo,))
    graficas = cur.fetchall()

    cur.execute("""SELECT m.fecharealizado,m.fechaincripcion,m.tipo,m.descripcion,t.cc_nit,t.nombre,t.fecharegistro,string_agg(nt.numero, ', ') AS telefonos
        FROM mantenimientoobservacionesugerencias  m
        JOIN tecnicos  t ON t.id  = m.idtecnico
        JOIN teletec  nt ON nt.idtecnicos  = m.idtecnico
        WHERE m.idequipo = %s
        GROUP BY 
        m.fecharealizado,
        m.fechaincripcion,
        m.tipo,
        m.descripcion,
        t.cc_nit,
        t.nombre,
        t.fecharegistro
    """, (id_equipo,))
    mantenimiento = cur.fetchall()

    cur.execute("""SELECT p.id,mp.marcas,p.serial,p.modelo,p.descripcion
        FROM perisfericos p
        JOIN marcaperis mp ON p.idmarca = mp.id
        WHERE p.idequipo = %s""", (id_equipo,))
    perisfericos = cur.fetchall()

    cur.execute("""SELECT mp.marca,p.pulgadas,p.voltaje,p.amperaje,p.serial
        FROM pantalla p
        JOIN marcapantalla mp ON p.idmarcapantalla = mp.id
        WHERE p.idequipo = %s""", (id_equipo,))
    pantallas = cur.fetchall()

    cur.execute("""SELECT id,nombresoftware,fechacompra,fechaexpiracion,clavelicencia
        FROM licenciasequipo  
        WHERE idequipo = %s""", (id_equipo,))
    licencias = cur.fetchall()

    cur.close()
    conn.close()

    # === Crear Excel ===
    wb = Workbook()
    ws = wb.active
    ws.title = "Equipo"

    # Estilos
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="4A90E2", end_color="4A90E2", fill_type="solid")
    center_align = Alignment(horizontal="center")

    def add_section(title, data_dict_or_list, headers=None):
        ws.append([title])
        ws.append([])

        if isinstance(data_dict_or_list, dict):
            for k, v in data_dict_or_list.items():
                val = v.strftime("%d/%m/%Y") if isinstance(v, datetime) else v
                ws.append([k, val if v else "No asignado"])
        elif isinstance(data_dict_or_list, list) and data_dict_or_list:
            ws.append(headers)
            for row in data_dict_or_list:
                ws.append([row.get(h, "No asignado") for h in headers])
            # aplicar estilos header
            for cell in ws[ws.max_row - len(data_dict_or_list) - 1]:
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = center_align
        ws.append([])
        ws.append([])

    # === Secciones ===
    add_section("Información General", equipo)
    add_section("Características Técnicas", caracteristicas)
    add_section("Almacenamiento", almacenamientos, ["marca", "capacidad", "tipo", "horas_uso", "estado"])
    add_section("Memoria RAM", memorias, ["marca", "tipo", "capacidad", "frecuencia"])
    add_section("Gráficas", graficas, ["tipo", "capacidad", "marca", "chipset"])
    add_section("Mantenimiento", mantenimiento, ["fecharealizado", "fechaincripcion", "tipo", "descripcion", "cc_nit", "nombre", "fecharegistro", "telefonos"])
    add_section("Periféricos", perisfericos, ["marcas", "serial", "modelo", "descripcion"])
    add_section("Pantallas", pantallas, ["marca", "pulgadas", "voltaje", "amperaje", "serial"])
    add_section("Licencias", licencias, ["nombresoftware", "fechacompra", "fechaexpiracion", "clavelicencia"])

    # === Guardar en buffer ===
    output = BytesIO()
    wb.save(output)
    output.seek(0)

    return send_file(
        output,
        as_attachment=True,
        download_name=f"equipo_{id_equipo}.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
